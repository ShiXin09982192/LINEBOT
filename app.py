import os, re, tempfile
from flask import Flask, request, abort
from linebot import LineBotApi, WebhookHandler
from linebot.models import MessageEvent, TextMessage, TextSendMessage
from docx import Document
from docx2pdf import convert
import boto3
from datetime import datetime

# ç’°å¢ƒè®Šæ•¸è¨­å®š
LINE_CHANNEL_TOKEN = os.getenv('LINE_CHANNEL_TOKEN')
LINE_CHANNEL_SECRET = os.getenv('LINE_CHANNEL_SECRET')
AWS_ACCESS_KEY = os.getenv('AWS_ACCESS_KEY')
AWS_SECRET_KEY = os.getenv('AWS_SECRET_KEY')
S3_BUCKET = os.getenv('S3_BUCKET')

line_bot_api = LineBotApi(LINE_CHANNEL_TOKEN)
handler = WebhookHandler(LINE_CHANNEL_SECRET)

s3 = boto3.client(
    's3',
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_KEY
)

app = Flask(__name__)

def parse_text(text):
    lines = text.splitlines()
    data = {'owner': '', 'address': '', 'items': []}
    for l in lines:
        if l.startswith('æ¥­ä¸»ï¼š'):
            data['owner'] = l.split('ï¼š', 1)[1].strip()
        elif l.startswith('åœ°å€ï¼š'):
            data['address'] = l.split('ï¼š', 1)[1].strip()
        m = re.match(r'([A-Za-z0-9]+)[\.\:\s]+(.+?)[\ï¼š\s]+([\d,]+)', l)
        if m:
            code, desc, amt = m.groups()
            amt = int(amt.replace(',', ''))
            data['items'].append((desc.strip(), amt))
    total = sum(amt for _, amt in data['items'])
    tax = int(total * 0.05)
    data.update({'total': total, 'tax': tax, 'grand_total': total + tax})
    return data

def fill_docx(data, out_docx_path):
    doc = Document('template.docx')
    table = doc.tables[0]

    # æ¸…é™¤é™¤äº†è¡¨é ­ä»¥å¤–çš„æ‰€æœ‰è¡Œ
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.text = ''

    # å¯«å…¥å“é …
    for i, (desc, amt) in enumerate(data['items'], start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = desc
        row[2].text = '1'
        row[3].text = 'å¼'
        row[4].text = f'{amt:,}'

    # ç¸½é‡‘é¡è³‡è¨Š
    doc.add_paragraph(f'åˆè¨ˆï¼š{data["total"]:,}')
    doc.add_paragraph(f'ç‡Ÿæ¥­ç¨…ï¼ˆ5%ï¼‰ï¼š{data["tax"]:,}')
    doc.add_paragraph(f'ç¸½è¨ˆï¼š{data["grand_total"]:,}')

    doc.save(out_docx_path)

def upload_and_get_url(file_path, key):
    s3.upload_file(file_path, S3_BUCKET, key, ExtraArgs={'ACL': 'public-read'})
    url = s3.generate_presigned_url(
        'get_object',
        Params={'Bucket': S3_BUCKET, 'Key': key},
        ExpiresIn=3600
    )
    return url

@app.route("/callback", methods=['POST'])
def callback():
    signature = request.headers.get('X-Line-Signature')
    body = request.get_data(as_text=True)

    try:
        handler.handle(body, signature)
    except Exception as e:
        abort(400)

    return 'OK'

@handler.add(MessageEvent, message=TextMessage)
def handle_message(event):
    text = event.message.text
    data = parse_text(text)

    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, f'å ±åƒ¹å–®_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx')
        pdf_path = docx_path.replace('.docx', '.pdf')

        fill_docx(data, docx_path)
        convert(docx_path, pdf_path)

        key_docx = os.path.basename(docx_path)
        key_pdf = os.path.basename(pdf_path)

        url_docx = upload_and_get_url(docx_path, key_docx)
        url_pdf = upload_and_get_url(pdf_path, key_pdf)

    msg = (f'å ±åƒ¹å–®å·²ç”Ÿæˆï¼š\n'
           f'ğŸ“„ PDFï¼š{url_pdf}\n'
           f'ğŸ“ Wordï¼š{url_docx}')

    line_bot_api.reply_message(event.reply_token, TextSendMessage(text=msg))

if __name__ == "__main__":
    app.run(port=8000)

